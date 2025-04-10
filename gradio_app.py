import gradio as gr
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline
import torch
from io import BytesIO
from docx import Document
from huggingface_hub import login
import os

# Load Mistral-7B model
def load_model():
    login(token="your_huggingface_token_here")
    tokenizer = AutoTokenizer.from_pretrained("mistralai/Mistral-7B-Instruct-v0.1")
    model = AutoModelForCausalLM.from_pretrained(
        "mistralai/Mistral-7B-Instruct-v0.1",
        torch_dtype=torch.float16,
        device_map="auto"
    )
    return pipeline("text-generation", model=model, tokenizer=tokenizer)

# Generate blog content
def generate_blog(topic, max_length=512, temperature=0.7, top_p=0.9):
    prompt = f"[INST] Write a detailed blog post about {topic}. [/INST]"
    result = generator(
        prompt, 
        max_length=max_length, 
        do_sample=True,
        temperature=temperature,
        top_p=top_p,
        truncation=True
    )[0]["generated_text"]
    return result

# Create Word document from blog content
def create_word_doc(content):
    doc = Document()
    doc.add_heading('Generated Blog Post', 0)
    doc.add_paragraph(content)
    
    # Save the document to a file instead of BytesIO
    file_path = "/tmp/generated_blog.docx"  # Modify this to save in a desired directory
    doc.save(file_path)
    return file_path

# Initialize Gradio interface
def blog_generator(topic, mode, max_length=512, temperature=0.7, top_p=0.9):
    global generator
    if 'generator' not in globals():
        generator = load_model()
    
    if mode == "Advanced":
        max_length = int(max_length)
        temperature = float(temperature)
        top_p = float(top_p)

    blog_content = generate_blog(topic, max_length, temperature, top_p)
    
    # Save blog content to Word document
    word_doc_path = create_word_doc(blog_content)
    
    return blog_content, word_doc_path  # Return the file path instead of the content of the BytesIO

# Gradio Interface
with gr.Blocks() as demo:
    gr.Markdown("# 📝 Mistral-Powered Blog Generator")
    
    with gr.Row():
        topic_input = gr.Textbox(label="Enter a blog topic:", value="How to Learn Machine Learning")
        mode_input = gr.Radio(["Simple", "Advanced"], label="Mode", value="Simple")
        with gr.Column():
            max_length_slider = gr.Slider(256, 1024, value=512, step=128, label="Maximum Length")
            temperature_slider = gr.Slider(0.1, 1.5, value=0.7, step=0.1, label="Temperature")
            top_p_slider = gr.Slider(0.1, 1.0, value=0.9, step=0.1, label="Top-p")
    
    output_blog = gr.Markdown()
    output_word = gr.File(label="Download Blog Post (Word Format)", type="filepath")
    
    generate_button = gr.Button("Generate Blog")
    
    generate_button.click(blog_generator, 
                          inputs=[topic_input, mode_input, max_length_slider, temperature_slider, top_p_slider],
                          outputs=[output_blog, output_word])

# Launch Gradio app
demo.launch(share=True)
