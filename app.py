import streamlit as st
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline
import torch
import warnings
from io import BytesIO
from docx import Document

# Suppress warnings
warnings.filterwarnings("ignore")

# Streamlit page configuration
st.set_page_config(
    page_title="Mistral Blog Generator",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for styling the app
st.markdown("""
    <style>
    .main .block-container {
        padding-top: 2rem;
    }
    .stButton>button {
        width: 100%;
    }
    </style>
    """, unsafe_allow_html=True)

st.title("📝 Mistral-Powered Blog Generator")

@st.cache_resource(show_spinner=True)
def load_model():
    with st.spinner("Loading Mistral-7B-Instruct model. This may take a few minutes..."):
        tokenizer = AutoTokenizer.from_pretrained("mistralai/Mistral-7B-Instruct-v0.1")
        model = AutoModelForCausalLM.from_pretrained(
            "mistralai/Mistral-7B-Instruct-v0.1",
            torch_dtype=torch.float16,
            device_map="auto"
        )
        return pipeline("text-generation", model=model, tokenizer=tokenizer)

def generate_blog(topic, max_length=512, temperature=0.7, top_p=0.9):
    prompt = f"[INST] Write a detailed blog post about {topic}. [/INST]"
    result = generator(
        prompt, 
        max_length=max_length, 
        do_sample=True,
        temperature=temperature,
        top_p=top_p
    )[0]["generated_text"]
    return result

def create_word_doc(content):
    """
    Creates a Word document from the generated blog content.
    """
    doc = Document()
    doc.add_heading('Generated Blog Post', 0)
    doc.add_paragraph(content)
    
    # Save the document to a BytesIO object
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Streamlit UI
def main():
    st.sidebar.title("⚙️ Generation Settings")
    mode = st.sidebar.radio("Mode", ["Simple", "Advanced"])
    
    max_length = 512
    temperature = 0.7
    top_p = 0.9
    
    if mode == "Advanced":
        max_length = st.sidebar.slider("Maximum Length", min_value=256, max_value=1024, value=512, step=128)
        temperature = st.sidebar.slider("Temperature", min_value=0.1, max_value=1.5, value=0.7, step=0.1)
        top_p = st.sidebar.slider("Top-p", min_value=0.1, max_value=1.0, value=0.9, step=0.1)
    
    with st.sidebar.expander("ℹ️ About Mistral-7B"):
        st.markdown("**Mistral-7B-Instruct** is a 7 billion parameter language model fine-tuned for following instructions.")
    
    topic = st.text_input("Enter a blog topic:", "How to Learn Machine Learning")
    
    if st.button("Generate Blog", type="primary"):
        if topic.strip():
            with st.spinner("Generating blog... please wait ⏳"):
                if 'generator' not in st.session_state:
                    st.session_state.generator = load_model()
                    
                generator = st.session_state.generator
                blog_content = generate_blog(topic, max_length, temperature, top_p)
                
                st.subheader(f"🧠 Generated Blog: {topic}")
                st.markdown(blog_content)
                
                # Create a download button for the Word document
                word_doc = create_word_doc(blog_content)
                st.download_button(
                    label="Download Blog Post (Word Format)",
                    data=word_doc,
                    file_name=f"{topic.replace(' ', '_').lower()}_blog.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Please enter a valid topic!")
    
    st.caption("Note: Blog generation may take 10-30 seconds depending on length and hardware.")

if __name__ == "__main__":
    main()
